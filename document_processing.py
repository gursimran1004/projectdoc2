from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import List

from docx import Document
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore
try:
    from PIL import Image
except Exception:
    Image = None  # type: ignore
try:
    import pytesseract
except Exception:
    pytesseract = None  # type: ignore
try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        PdfReader = None  # type: ignore


@dataclass
class DocumentSection:
    text: str
    source_label: str


@dataclass
class ParsedDocument:
    name: str
    extension: str
    text: str
    sections: List[DocumentSection]


def _normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _fix_fragmented_ocr_words(text: str) -> str:
    # Join OCR-fragmented words like "c o n s i s t i n g" -> "consisting".
    pattern = re.compile(r"\b(?:[A-Za-z]\s+){3,}[A-Za-z]{1,3}\b")

    def _join(match: re.Match) -> str:
        return match.group(0).replace(" ", "")

    fixed = pattern.sub(_join, text)

    # Merge noisy short-token runs such as "pro ce du re" -> "procedure".
    tokens = re.findall(r"[A-Za-z]+|[^A-Za-z]+", fixed)
    merged: List[str] = []
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.isalpha() and len(tok) <= 2:
            j = i
            run: List[str] = []
            while j < len(tokens) and tokens[j].isalpha() and len(tokens[j]) <= 2:
                run.append(tokens[j])
                j += 1
                if j < len(tokens) and tokens[j].isspace():
                    j += 1
            if len(run) >= 3 and sum(len(x) for x in run) >= 6:
                merged.append("".join(run))
                i = j
                continue
        merged.append(tok)
        i += 1

    return _normalize_text("".join(merged))


def _ocr_image_to_text(image) -> str:
    """Extract text from image using Tesseract OCR with preprocessing."""
    if pytesseract is None:
        return ""
    try:
        # Convert to grayscale
        gray = image.convert("L")
        # Apply threshold to create binary image (helps OCR)
        bw = gray.point(lambda x: 0 if x < 150 else 255, "1")
        # Use Tesseract with optimized config for scanned documents
        text = pytesseract.image_to_string(bw, config="--oem 3 --psm 6 --dpi 300")
        if not text.strip():
            # Fallback: try without preprocessing
            text = pytesseract.image_to_string(image, config="--oem 3 --psm 6")
    except Exception as e:
        print(f"OCR error: {e}")
        try:
            text = pytesseract.image_to_string(image)
        except Exception:
            return ""
    
    return _fix_fragmented_ocr_words(text)


def _low_quality_text(text: str) -> bool:
    """Detect if extracted text is of low quality (likely needs OCR)."""
    if not text.strip():
        return True
    words = text.split()
    if not words:
        return True
    
    # Check for single character words (indicates extraction failure)
    single_char = sum(1 for w in words if len(w) == 1)
    ratio = single_char / max(1, len(words))
    
    # Check for excessive whitespace or repeated characters
    whitespace_ratio = sum(1 for c in text if c.isspace()) / max(1, len(text))
    repeated_chars = sum(1 for i in range(len(text) - 1) if text[i] == text[i+1]) / max(1, len(text) - 1)
    
    # Lowered threshold from 0.30 to 0.20 for better OCR triggering
    return ratio > 0.20 or whitespace_ratio > 0.50 or repeated_chars > 0.40


def _extract_pdf_text_pymupdf(file_bytes: bytes, force_ocr: bool = False) -> tuple[str, List[DocumentSection]]:
    """Extract text from PDF using PyMuPDF with optional OCR fallback."""
    if fitz is None:
        return "", []
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        print(f"Error opening PDF: {e}")
        return "", []
    
    sections: List[DocumentSection] = []
    pages_text: List[str] = []
    
    for idx, page in enumerate(doc, start=1):
        try:
            block_items = page.get_text("blocks", sort=True)
            block_text = " ".join(item[4] for item in block_items if len(item) > 4 and str(item[4]).strip())
            page_text = _normalize_text(block_text)
            
            # Determine if OCR is needed
            needs_ocr = force_ocr or _low_quality_text(page_text)
            
            if needs_ocr and Image is not None and pytesseract is not None:
                try:
                    # Render page at higher DPI for better OCR accuracy
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5))
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_text = _ocr_image_to_text(img)
                    
                    # Use OCR text if it's better quality
                    if ocr_text and len(ocr_text.split()) > len(page_text.split()) * 0.8:
                        page_text = ocr_text
                except Exception as e:
                    print(f"OCR failed for page {idx}: {e}")
                    # Keep original extraction if OCR fails
                    pass
            
            if page_text:
                pages_text.append(page_text)
                sections.append(DocumentSection(text=page_text, source_label=f"Page {idx}"))
        
        except Exception as e:
            print(f"Error processing page {idx}: {e}")
            continue
    
    return _normalize_text("\n".join(pages_text)), sections


def _read_pdf(file_bytes: bytes, force_ocr: bool = False) -> tuple[str, List[DocumentSection]]:
    """Read PDF with fallback options."""
    # Prefer PyMuPDF when available for better extraction quality
    if fitz is not None:
        text, sections = _extract_pdf_text_pymupdf(file_bytes, force_ocr=force_ocr)
        if text:
            return text, sections

    if PdfReader is None:
        raise ImportError(
            "PDF support requires 'pypdf' (or 'PyPDF2'). Install dependencies using: pip install -r requirements.txt"
        )
    
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot read PDF: {e}")
    
    pages: List[str] = []
    sections: List[DocumentSection] = []
    
    for idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = _normalize_text(page.extract_text() or "")
            if page_text:
                pages.append(page_text)
                sections.append(DocumentSection(text=page_text, source_label=f"Page {idx}"))
        except Exception as e:
            print(f"Error extracting text from page {idx}: {e}")
            continue
    
    if not pages:
        raise ValueError("Could not extract text from any pages in the PDF.")
    
    return _normalize_text("\n".join(pages)), sections


def _read_docx(file_bytes: bytes) -> tuple[str, List[DocumentSection]]:
    """Extract text from DOCX files."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot read DOCX: {e}")
    
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    sections = [
        DocumentSection(text=paragraph, source_label=f"Paragraph {idx}")
        for idx, paragraph in enumerate(parts, start=1)
    ]
    return _normalize_text("\n".join(parts)), sections


def _read_txt(file_bytes: bytes) -> tuple[str, List[DocumentSection]]:
    """Extract text from TXT files."""
    try:
        text = _normalize_text(file_bytes.decode("utf-8", errors="ignore"))
    except Exception as e:
        raise ValueError(f"Cannot read TXT: {e}")
    
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections = [DocumentSection(text=line, source_label=f"Line {idx}") for idx, line in enumerate(lines, start=1)]
    return text, sections


def _read_image(file_bytes: bytes) -> tuple[str, List[DocumentSection]]:
    """Extract text from image files using OCR."""
    if Image is None or pytesseract is None:
        raise ImportError(
            "Image scanning requires 'Pillow' and 'pytesseract'. Install with: pip install pillow pytesseract"
        )
    
    try:
        image = Image.open(io.BytesIO(file_bytes))
        text = _ocr_image_to_text(image).strip()
    except Exception as e:
        raise ValueError(f"Cannot process image: {e}")
    
    sections = [DocumentSection(text=text, source_label="Image OCR")] if text else []
    return text, sections


def parse_uploaded_file(uploaded_file, force_ocr: bool = False) -> ParsedDocument:
    """Parse uploaded file and extract text."""
    file_name = uploaded_file.name
    extension = file_name.lower().split(".")[-1]
    file_bytes = uploaded_file.read()

    try:
        if extension == "pdf":
            text, sections = _read_pdf(file_bytes, force_ocr=force_ocr)
        elif extension == "docx":
            text, sections = _read_docx(file_bytes)
        elif extension == "txt":
            text, sections = _read_txt(file_bytes)
        elif extension in {"png", "jpg", "jpeg"}:
            text, sections = _read_image(file_bytes)
        else:
            raise ValueError("Unsupported file type. Please upload PDF, DOCX, TXT, PNG, JPG, or JPEG.")

        if not text:
            raise ValueError("Could not extract text from the uploaded file. Try enabling 'Force OCR' in Settings.")

        return ParsedDocument(name=file_name, extension=extension, text=text, sections=sections)
    
    except Exception as e:
        raise ValueError(f"Error processing {file_name}: {str(e)}")
